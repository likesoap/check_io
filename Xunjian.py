import re
import os
import docx
import datetime
import sys
import csv
import glob
import subprocess
import arial10
import xlwt
import shutil
import time
from docx.shared import Pt,Inches
from multiprocessing import Pool,freeze_support
from itertools import zip_longest
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QApplication,QGridLayout,QCheckBox,QPushButton,QWidget

##Get the last level path
getpath = os.path.dirname(os.getcwd())
###Create a "auto-crt.bat"
batfile = os.path.join(getpath, "auto-crt.bat")
path = os.getcwd()

'''
     Three steps here:
     1. Write and run a bat which will automatically open "SecureCRT" and run the crt Script
     2. Filter some key words
     3. Write the key words to different format, e.g: csv, xls, doc
     The final report will be named as "巡检报告.xls"
     '''


def getbat(batfile): ##Get an auto-crt.bat
    ####Please note that you need to change the crt name if you rename the file
    ####Change "Xunjian-ct.py" to what you have renamed to

    with open(batfile,'w+') as batman:
        batman.write("@echo off")
        batman.write("\n")
        batman.write("start {path}\App\SecureCRT\SecureCRT.exe /SCRIPT {path}\Xunjian-crt.py".format(path=getpath))
        batman.write("\n")
        batman.write(":end")
        batman.write("\n")
        batman.write("exit")
    batman.close()


class Window(QWidget):
    def __init__(self,parent=None):
        super(Window,self).__init__(parent)
        self.initUI()

    def initUI(self):
        self.listCheckBox=["系统时间", "设备名称", "设备版本", "设备类型", "运行时间","cpu利用率"
                             ,'内存使用','板卡信息','板卡切换情况','邻居信息','温度与电源信息','ACL','HomelessVideo']

        grid=QGridLayout()

        self.resize(300, 200)
        self.setWindowTitle("来快活呀")

        for i,v in enumerate(self.listCheckBox):
            self.listCheckBox[i]=QCheckBox(v)
            grid.addWidget(self.listCheckBox[i],i,0)

        self.checkbox=QCheckBox("选择全部")
        self.checkbox.stateChanged.connect(self.selectall)

        self.button=QPushButton("Run")
        self.button.clicked.connect(self.buttonclicked)

        grid.addWidget(self.checkbox,25,0,1,2)
        grid.addWidget(self.button, 26,0,1,2)
        self.setLayout(grid)


    def buttonclicked(self):
        ###It will return a varlist.
        ###It will display what you have chosen.
        varlist=[]
        for i,v in enumerate(self.listCheckBox):
            if v.checkState():
                varlist.append(v.text())
        print(varlist)
        ###After a varlist is chosen, run the function in order to export the chosen items
        t1 = time.time()
        readwritevalue(varlist)
        GetFormated()  ##Get the Final output of Xunjian-Report.csv
        Getxlsx()  ##Get the Xunjian-Report.xls
        GetDocx()  ##Get the docx format
        print("Run time %s" % (time.time() - t1))

        ###Multiple process###
        ###Don't know how to proceed with multiprocessing under the PyQt yet.
        ##The following will spend more time to complete the programme.
        ##From the url https://stackoverflow.com/questions/15675043/multiprocessing-and-gui-updating-qprocess-or-multiprocessing
        ##It describes how to use multiprocessing and QThread together.

        # t1 = time.time()
        # pool = Pool(4)
        # pool.map(readwritevalue,  varlist)
        # GetFormated()  ##Get the Final output of Xunjian-Report.csv
        # Getxlsx()  ##Get the Xunjian-Report.xls
        # GetDocx()  ##Get the docx format
        # pool.close()
        # pool.join()
        # print("Run time %s" %(time.time()-t1))

    def selectall(self):
        if self.checkbox.checkState()==Qt.Checked:
            for i,v in enumerate(self.listCheckBox):
                self.listCheckBox[i].setChecked(True)
        else:
            for i,v in enumerate(self.listCheckBox):
                self.listCheckBox[i].setChecked(False)

def readwritevalue(varlist):
    ###Three parts
    ###Filter the key words from the log files.
    ###Matching the varlist, export the key words to a ".csv" file
    ###Matching the varlist, modify the "Format-1.csv" to a new "Format-after.csv" file

    for file in gettxtpath():
        ##Get the host name from the filename. Please note that the file now contains the whole path and the extension
        host = file.split("\\")[-1].split(".")[0]
        w0 = host + "show video homeless-streams"
        w1 = host + "#show ip ospf neighbor"
        w2 = host + "#show ip bgp summary"
        w3 = host + "#show isis neighbor"
        w4 = host + "#show cable modem summary total"
        w5 = host + "#show envm"
        w6 = host + "#show ha log"
        w7 = host + "#show log"
        Module = []
        Module_Uptime = []
        Power = []
        Temperature = []
        # Startup = []
        # Router = []
        ACL = []

        filename = os.path.split(file)[1]  ##Get the filename including the extension
        file1 = filename + ".csv"  ##The file will be like "161.txt.csv"

        ###Change the folder so that the csv,xlsx can be under Xunjian folder
        path1 = os.path.join(getpath, 'Xunjian')
        file1path = os.path.join(path1, file1)

        ##Filter the neighbor data and the environment data
        with open(file, 'r', encoding='UTF-8') as fp_nei:
            fp_reader = fp_nei.read()
            # fp_reader=fp_reader.replace("\n",'/r/n')
            pat = re.compile(w0 + '(.*?)' + w1, re.S)
            homelessvideo = pat.findall(fp_reader)
            pat = re.compile(w1 + '(.*?)' + w2, re.S)
            ospfnei = pat.findall(fp_reader)
            pat = re.compile(w2 + '(.*?)' + w3, re.S)
            bgpnei = pat.findall(fp_reader)
            pat = re.compile(w3 + '(.*?)' + w4, re.S)
            isisnei = pat.findall(fp_reader)
            pat = re.compile(w5 + '(.*?)' + w6, re.S)
            envm = pat.findall(fp_reader)
            pat = re.compile(w6 + '(.*?)' + w7, re.S)
            showhalog = pat.findall(fp_reader)
        fp_nei.close()
        halog = showhalog
        neiarray = [ospfnei, bgpnei, isisnei]  ##Get the neighbor list
        array = [envm, Power, Temperature]


        ###For the remaining data
        with open(file, 'r', encoding='UTF-8') as fp:
            for line in fp:
                line_stream = line.strip().replace('\x00', '')
                if "System Time" in line:
                    systemtime = line_stream.split(":")[1]

                if "show version" in line:
                    hostname = line.split('#')[0]

                ##########Check the version##############
                if "Running Image" in line:
                    version = line_stream.split(",")[0:3]
                    version = ''.join(version)

                #########Check the product############
                if "Product" in line:
                    product = line.split(':')[1].split(',')[0]

                ##########Check the uptime and the product type#############
                if "System Uptime" in line:
                    uptime = line.split()[2:]
                    uptime = ''.join(uptime)
                    uptime = str(uptime).replace(",", ":")

                ############Check the CPU#############################
                if "CPU states" in line:
                    # print(line_stream)
                    cpu = line_stream.split(",")[3].strip()
                    cpu = re.findall("[0-9]+\.[0-9]+", cpu)

                ##########Check the Memory############################
                if "MemTotal" in line:
                    MemTotal = re.findall("\d+", line_stream)
                    MemTotal = int(MemTotal[0])

                if "MemFree" in line_stream:
                    MemFree = re.findall("\d+", line_stream)
                    MemFree = int(MemFree[0])

                    ############Check the running state of module#########
                if "Module" in line or "Uptime" in line:
                    Module += re.findall("^Module \d+.+$", line_stream)
                    Module_Uptime += re.findall("^Uptime.+$", line_stream)

                    #########Check the power###############
                if "power module" in line:
                    line_stream = line_stream.strip("[").replace("]", "")
                    Power.append(line_stream)

                    #########Check temperature##########
                if "Module temperature is high" in line or "overheating" in line:
                    Temperature.append(line_stream)

                    #########Check ACL###################
                if "sunrpc" in line:
                    ACL.append(line_stream)
                    return ACL
                    #######When detecting the "show tech", skip the loop##########
                if "show tech" in line:
                    break

                ####In the future, if you want to add something ######
                ####Please add it from this line ##################
                ####e.g #############################
                ###Check the log###########################
                # if "failed to become startup" in line:
                #     line_stream=line_stream.replace(",", " and will")
                #     #line_stream=line_stream.strip("[").replace("]", "").replace(",", " and will")
                #     Startup.append(line_stream)
                # if "-WA-ROUTER-0" in line:
                #     line_stream=line_stream.replace(",", " and")
                #     #line_stream = line_stream.strip("[").replace("]", "").replace(",", " and")
                #     Router.append(line_stream)


         ##########################################################
         ##Step3: Print out the result
        # The content in the "Format.csv" will be like
        # System Clock
        # Hostname
        # SMM  Version
        # Product
        # System Uptime
        # CPU status
        # MemFree/MemTotal
        # Module Status
        # OSPF Neighbor
        # BGP Neighbor
        # ISIS Neighbor
        # Power, temperature status
        # Power module alert
        # Temperature alert
        # ACL

        # According to what you have chosen from the checkbox (class Windows()), the filtered data will be printed with ".csv" format.
        # We also read the "Format-1.csv"(Chinese) to write a new format csv in order to merge the filtered data in the end.
        # You will be able to have a report which you can chose what to print out.

        ##Please note that you will need to modify the list and outputlist these two list and add what you want to append in the end of the list.
        ##Right now, the order is strictly restricted. You cannot change the order of the list.

         ##########################################################

        with open(file1path, 'w+', encoding='UTF-8', newline='') as fp1:
            fp1.write("Result\n")
            w = csv.writer(fp1)

            ###CPU statistics
            cpu = cpu[0]
            if cpu >= str(70):
                cpu="CPU utilization is normal with idle " + cpu
            else:
                cpu = "CPU utilization is abnormal with idle " + cpu

            ###Memory statistics
            Mem = MemFree/MemTotal
            if Mem > 0.45:
                Mem = "Memory is with good with %.2f%%" %(Mem * 100)
            else:
                Mem="Memory is with abnormal with %.2f%%" %(Mem * 100)

            ##Filter the module type and uptime
            c=[]
            for i in range(len(Module)):
                if Module_Uptime != []:
                    c.append(Module[i] + "\n")
                    c.append(Module_Uptime[i] + "\n")
                else:
                    c.append(Module[i] + "\n")
            c_1 = ''.join(c)
            Module_F = [c_1]


            ######################################################################
            ##Please note that you will need to modify the list and outputlist these two lists and add what you want to append in the end of the lists.
            ##Right now, the order is strictly restricted. You cannot change the order of the list.
            #######################################################################
            ### It is used to match the item in varlist
            list = ["系统时间", "设备名称", "设备版本", "设备类型", "运行时间","cpu利用率","内存使用"]
            ###"outputlist' is used to export the value of the items
            outputlist = [systemtime, hostname, version, product, uptime, cpu, Mem]

            ###If the item in the list matches the item in the varlist, print and export to csv file
            ###The order of the list and varlist should be the same.

            for j in range(len(list)):
                if list[j] in varlist:
                    print(outputlist[j])
                    print(outputlist[j],file=fp1)

            ##"Module_F", "neiarray" and "array", "ACL" need to individually specify as each item is a list.

            ##For the module and uptime
            if "板卡信息" in varlist:
                w.writerow(Module_F)

            if "板卡切换情况" in varlist:
                if "UPS" in halog or "QAM" in halog or "SMM" in halog:
                    print(halog)
                    w.writerow(halog)
                else:
                    print("'show ha log' is not supported")
                    w.writerow(["'show ha log' is not supported"])

            ##Write the neighbor list including ospf, bgp, isis
            if "邻居信息" in varlist:
                for i in range(len(neiarray)):
                    tip = ['Not configured']
                    if neiarray[i] != ['\n']:
                        w.writerow(neiarray[i])
                        # print(neiarray[i])
                    else:
                        w.writerow(tip)
                        # print(tip)

            ##For the other logs like Power and temperature....
            if "温度与电源信息" in varlist:
                for id in range(len(array)):
                    if array[id] == []:
                        print("Good")
                        w.writerow(["Good"])
                    else:
                        # print(array[id])
                        out = ','.join(array[id])
                        print(out)
                        w.writerow([out])

            ##For ACL ###########################################
            ##You can also add more lines to check any improvement on the configuration.
            if "ACL" in varlist:
                if ACL == []:
                    # print("No sunrpc ACL is detected. Please implement the ACL")
                    print("No sunrpc ACL is detected. Please implement the ACL", file=fp1)
                else:
                    # print(ACL)
                    print(ACL, file=fp1)

            if "HomelessVideo" in varlist:
                print(homelessvideo)
                # w.writerow([homelessvideo])

    ######Read the "Format-1.csv" format, and write into a new file using selected items.

    with open ("Format-1.csv") as fin, open("Format-after.csv",'w',newline='') as fout:
        reader=csv.reader(fin)
        writer=csv.writer(fout)
        rows = [row for row in reader]

        ###Write the first row (The subject of the file)
        ###Please note that the id starts from zero. It means the row 1 in the csv file.

        writer.writerow(rows[0])

        ###The same list on the above
        list = ["系统时间", "设备名称", "设备版本", "设备类型", "运行时间","cpu利用率",'内存使用', '板卡信息']

        ##Compared with the "Format" csv, check the row number. Reduce one if the id you need to input
        for j in range(len(list)):
            if list[j] in varlist:
                list[j]= rows[j+1]
                writer.writerow(list[j])

        if '板卡切换情况' in varlist:
            # print(rows[9])
            writer.writerow(rows[9])

        ##For the row 10 to row 12
        if '邻居信息' in varlist:
            for i in range(10,13):
                print(rows[i])
                writer.writerow(rows[i])
        ##For the row 13 to row 15
        if '温度与电源信息' in varlist:
            for r in range(13,16):
                print(rows[r])
                writer.writerow(rows[r])
        ##For the row 16
        if 'ACL' in varlist:
            print(rows[16])
            writer.writerow(rows[16])
        ##You can add more lines from here##################

        if "Homeless Video" in varlist:
            writer.writerow(rows[17])

def gettxtpath():  ##'.txt'.
    #Get all the filename with txt format. It is for data filtering
    ##The filename is under the Log folder
    Logpath=os.path.join(getpath, 'Log')
    file_path=[]
    f_list=os.listdir(Logpath)
    for file in f_list:
        #os.path.splitext()
        if os.path.splitext(file)[1] == '.txt':
            filepath=os.path.join(Logpath,file)
            file_path.append(filepath)
    print(file_path)
    return file_path


def getcsvName():
    #Get all the filename with txt.csv format..
    file_list=[]
    csv_list=[]
    f_list=os.listdir(path)
    for file in f_list:
        #os.path.splitext()
        if os.path.splitext(file)[1] == '.csv':
            file_list.append(file)
    for file in file_list:
        if file.endswith(".txt.csv"):   ###It has contained the format
            csv_list.append(file)
    return csv_list    ##It has a list for the ".txt.csv" file like "161.txt.csv"


def CsvMerge():
    ##Merge the file which contains key word with the "Format-after.csv" file
    for file in getcsvName():
        file1=file+'.csv'
        formatfile='Format-after.csv'   ###The re-defined format
        with open(file1,'w',newline='') as fout,open(formatfile) as fin, open(file) as fin1:
            reader=csv.reader(fin)
            reader1=csv.reader(fin1)
            writer = csv.writer(fout, delimiter=',', lineterminator='\n')
            for line1,line2 in zip_longest(reader,reader1):
                writer.writerow(line1+line2)

def GetFormated():
    CsvMerge()
    getResult()  ##Get the Xunjian-Report.csv

def getResult():
    ##Get the Final output of Xunjian-Report.csv. Merged all the csv with txt.csv.csv
    ##Use glob.glob to get all the "txt.csv.csv" files.
    allfiles = glob.glob(path + "/*.txt.csv.csv")
    df_out_filename='Xunjian-Report.csv'
    #write_headers=True
    with open(df_out_filename,'w+',newline='') as fout:
        writer=csv.writer(fout)
        for filename in allfiles:
            with open(filename) as fin:
                reader=csv.reader(fin)
                #headers=next(reader)
                #if write_headers:
                #    write_headers=False     # Only write headers once.
                    #writer.writerow(headers)
                writer.writerows(reader)
                writer.writerow([])


def Getxlsx():
    '''
    Read the csv and save it with xls format
    Also Set the format of the cells
    ##Wrap at right
    ##Horz_left
    ##Vert_center
    '''
    myexcel=xlwt.Workbook()
    ##This is to modify columns to max size of any entry.
    mysheet=FitSheetWrapper(myexcel.add_sheet("Xunjian"))  ##Apply the class
    ##Set the format of the cells
    al1=xlwt.Alignment()
    al1.wrap=xlwt.Alignment.WRAP_AT_RIGHT
    al1.horz=xlwt.Alignment.HORZ_LEFT
    al1.vert=xlwt.Alignment.VERT_CENTER
    style=xlwt.XFStyle()
    style.alignment=al1
    ##Set the bold font
    # font=xlwt.Font()
    # font.bold=True
    # style1=xlwt.XFStyle()
    # style1.font=font

    ###Read the csv and save it with xls format
    csvfile=open("Xunjian-Report.csv","r")
    reader=csv.reader(csvfile)
    writevaluexlsx(reader,mysheet,style)
    csvfile.close()

    ##Separate with different sheets.
    allfiles = glob.glob(os.getcwd() + "/*.txt.csv.csv")
    for file in allfiles:
        ##Get the host name and name it as the sheet name
        hostexten=os.path.split(file)[1]
        host=os.path.splitext(hostexten)[0].split(".")[0]
        with open(file) as fin:
            reader = csv.reader(fin)
            newsheet = FitSheetWrapper(myexcel.add_sheet(host))
            writevaluexlsx(reader,newsheet,style)
        fin.close()

    ##Save the excel
    myexcel.save("巡检报告.xls")
    ###Copy the file to upper folder which is is in the same path of the crt.py
    shutil.copyfile("巡检报告.xls","{path}/巡检报告.xls".format(path=getpath))


def writevaluexlsx(reader,sheet,style):
    ##Read the csv file and write it to the xlsx
    l=0
    for line in reader:
        r=0
        for i in line:
            sheet.write(l,r,i,style)
            r=r+1
        l=l+1

class FitSheetWrapper(object):
    #https://stackoverflow.com/questions/6929115/python-xlwt-accessing-existing-cell-content-auto-adjust-column-width
    """Try to fit columns to max size of any entry.
    To use, wrap this around a worksheet returned from the
    workbook's add_sheet method, like follows:
        sheet = FitSheetWrapper(book.add_sheet(sheet_name))
    The worksheet interface remains the same: this is a drop-in wrapper
    for auto-sizing columns.
    """
    def __init__(self, sheet):
        self.sheet = sheet
        self.widths = dict()

    def write(self, r, c, label='', *args, **kwargs):
        self.sheet.write(r, c, label, *args, **kwargs)
        width = int(arial10.fitwidth(label))

        if width > self.widths.get(c, 0):
            self.widths[c] = width
            self.sheet.col(c).width = width

    def __getattr__(self, attr):
        return getattr(self.sheet, attr)


def GetDocx():
    ##Get docx
    ###To generate the docx

    doc = docx.Document(docx=os.path.join(os.getcwd(), 'default.docx'))
    date = datetime.datetime.now()

    font = doc.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(9)

    doc.add_heading('巡检报告', level=0)
    doc.add_heading('Date: %s/%s/%s' % (date.day, date.month, date.year), level=1)
    ##Add a blank row
    doc.add_paragraph("\n")

    allfiles = glob.glob(path + "/*.txt.csv.csv")

    for filename in allfiles:
        # Get the host name
        host = filename.split("\\")[-1].split(".")[0]
        with open(filename) as fin:
            csv_reader = csv.reader(fin)
            csv_headers = next(csv_reader)
            csv_cols = len(csv_headers)

            ##Add the hostname of the system
            doc.add_heading('System Name: %s' % host, level=1)
            doc.add_paragraph("\n")
            ##Add the table
            table = doc.add_table(rows=1, cols=csv_cols)
            table.autofix = True
            set_col_widths(table)  ##Set the width

            ##For the table aubject
            hdr_cells = table.rows[0].cells
            ##The subject content
            for i in range(csv_cols):
                hdr_cells[i].text = csv_headers[i]
            ##For other content
            for row in csv_reader:
                row_cells = table.add_row().cells
                for i in range(len(row)):
                    row_cells[i].text = row[i]

            doc.add_page_break()

    # paragraph.add_run('xxx').bold=True
    doc.save("Xunjian-Report.docx")
    ###Copy the file to upper folder which is is in the same path of the crt.py
    shutil.copyfile("Xunjian-Report.docx", "{path}/Xunjian-Report.docx".format(path=getpath))

def set_col_widths(table):
    ##Modify the table widths of the docx
    # https://stackoverflow.com/questions/43051462/python-docx-how-to-set-cell-width-in-tables

    widths=(Inches(0.3),Inches(0.5),Inches(1.2),Inches(2))
    for row in table.rows:
        for (idx,width) in enumerate(widths):
            row.cells[idx].width=width


if __name__ == "__main__":

    ##Add support for when a program which uses multiprocessing has been frozen to produce a Windows executable. (Has been tested with py2exe, PyInstaller and cx_Freeze.)
    #One needs to call this function straight after the if __name__ == '__main__' line of the main module.
    freeze_support()

    ##########################################################
    ##Step1: Generate a bat and run it. After you manuall click the "confirm" buttom in the crt, it will quit.
    ##Then we will move to the step2
    ##########################################################

    # getbat(batfile)  ##Write a "auto-crt.bat"
    # child = subprocess.Popen("auto-crt.bat", cwd=getpath, shell=True, stdin=subprocess.PIPE, stdout=subprocess.PIPE,
    #                          stderr=subprocess.PIPE)
    # stdout, stderr = child.communicate()
    # print(child.returncode)
    #
    # #########################################################
    # #Step2-Step3: Filter and write.
    # #########root################################################

    # t1 = time.time()
    app = QApplication(sys.argv)
    clock = Window()
    clock.show()
    sys.exit(app.exec_())

